import logging
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import REFINED_DIR

logger = logging.getLogger("oral-history-tools")


def export_txt(refined_text, output_dir=None, filename=None):
    """Export refined text to a .txt file.

    Args:
        refined_text: The full refined transcript string.
        output_dir: Output directory (default: REFINED_DIR).
        filename: Output filename without extension (default: timestamped).

    Returns:
        Path to the created file.
    """
    output_dir = output_dir or REFINED_DIR
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    path = os.path.join(output_dir, f"{filename}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(refined_text)

    logger.info("Exported TXT: %s (%d chars)", path, len(refined_text))
    return path


def export_docx(refined_text, output_dir=None, filename=None, title=None, metadata=None):
    """Export refined text to a formatted .docx file.

    Args:
        refined_text: The full refined transcript string.
        output_dir: Output directory (default: REFINED_DIR).
        filename: Output filename without extension (default: timestamped).
        title: Document title for the title page.
        metadata: Optional dict with stats to include on title page
                  (e.g. {"pages": 12, "source_files": 3, "tokens_used": 5000}).

    Returns:
        Path to the created file.
    """
    output_dir = output_dir or REFINED_DIR
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"transcript_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title page
    _add_title_page(doc, title or "Document Transcription", metadata)

    # Body — split on page markers and double newlines
    _add_body(doc, refined_text)

    path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(path)

    logger.info("Exported DOCX: %s (%d chars)", path, len(refined_text))
    return path


def _add_title_page(doc, title, metadata):
    """Add a formatted title page."""
    # Title
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Transcribed: {datetime.now().strftime('%B %d, %Y')}")
    date_run.font.size = Pt(12)
    date_run.italic = True

    # Metadata summary
    if metadata:
        doc.add_paragraph()  # spacer
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = []
        if "source_files" in metadata:
            lines.append(f"Source files: {metadata['source_files']}")
        if "pages" in metadata:
            lines.append(f"Pages transcribed: {metadata['pages']}")
        if "tokens_used" in metadata:
            lines.append(f"API tokens used: {metadata['tokens_used']:,}")
        meta_run = meta_para.add_run("\n".join(lines))
        meta_run.font.size = Pt(10)
        meta_run.font.color.rgb = None  # default color

    # Page break after title page
    doc.add_page_break()


def _add_body(doc, text):
    """Add the transcript body, handling page markers as section headers."""
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue

        # Detect page marker lines: --- PAGE N: file, page N ---
        if block.startswith("--- PAGE") and block.endswith("---"):
            para = doc.add_paragraph()
            run = para.add_run(block)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = None
            para.space_before = Pt(18)
            para.space_after = Pt(6)
        else:
            doc.add_paragraph(block)


def export_raw_txt(raw_results, output_dir=None, filename=None):
    """Export raw page-by-page transcription results to a .txt file.

    Args:
        raw_results: List of page result dicts from transcription (must have
            sequence, source_file, source_page, text, status keys).
        output_dir: Output directory (default: REFINED_DIR).
        filename: Output filename without extension (default: timestamped).

    Returns:
        Path to the created file.
    """
    output_dir = output_dir or REFINED_DIR
    os.makedirs(output_dir, exist_ok=True)

    if not filename:
        filename = f"transcript_raw_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    parts = []
    for r in raw_results:
        if r.get("status") != "ok":
            continue
        header = f"--- PAGE {r['sequence']}: {r['source_file']}, page {r['source_page']} ---"
        parts.append(f"{header}\n{r.get('text', '')}")

    text = "\n\n".join(parts)
    path = os.path.join(output_dir, f"{filename}.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

    logger.info("Exported raw TXT: %s (%d chars, %d pages)", path, len(text), len(parts))
    return path


def export_all(refined_text, output_dir=None, filename=None, title=None, metadata=None,
               raw_results=None):
    """Export to .txt, .docx, and optionally raw .txt. Returns dict of paths."""
    txt_path = export_txt(refined_text, output_dir, filename)
    docx_path = export_docx(refined_text, output_dir, filename, title, metadata)
    result = {"txt": txt_path, "docx": docx_path}

    if raw_results:
        raw_filename = f"{filename}_raw" if filename else None
        raw_path = export_raw_txt(raw_results, output_dir, raw_filename)
        result["raw"] = raw_path
        print(f"\nExported:")
        print(f"  RAW:  {raw_path}")
    else:
        print(f"\nExported:")

    print(f"  TXT:  {txt_path}")
    print(f"  DOCX: {docx_path}")
    return result
