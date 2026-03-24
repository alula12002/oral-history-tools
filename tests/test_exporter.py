import os
import sys
import tempfile
import unittest

from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from shared.exporter import export_txt, export_docx, export_all, export_raw_txt


SAMPLE_TEXT = """--- PAGE 1: notebook.pdf, page 1 ---
The house was built in 1890. My grandmother [illegible] told me about it.

She said "we ain't never had nothing like it before." [guess?]

--- PAGE 2: notebook.pdf, page 2 ---
The family moved north in 1923. [STAMP: Official County Record]

There were seven children, all born in that same house."""


class TestExportTxt(unittest.TestCase):

    def test_creates_txt_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_txt(SAMPLE_TEXT, output_dir=tmpdir, filename="test_out")
            self.assertTrue(os.path.exists(path))
            self.assertTrue(path.endswith(".txt"))
            with open(path, encoding="utf-8") as f:
                content = f.read()
            self.assertEqual(content, SAMPLE_TEXT)

    def test_default_filename_is_timestamped(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_txt("hello", output_dir=tmpdir)
            basename = os.path.basename(path)
            self.assertTrue(basename.startswith("transcript_"))

    def test_creates_output_dir_if_missing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            nested = os.path.join(tmpdir, "sub", "dir")
            path = export_txt("test", output_dir=nested, filename="out")
            self.assertTrue(os.path.exists(path))

    def test_preserves_markers(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_txt(SAMPLE_TEXT, output_dir=tmpdir, filename="markers")
            with open(path, encoding="utf-8") as f:
                content = f.read()
            self.assertIn("[illegible]", content)
            self.assertIn("[guess?]", content)
            self.assertIn("[STAMP: Official County Record]", content)


class TestExportDocx(unittest.TestCase):

    def test_creates_docx_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_docx(SAMPLE_TEXT, output_dir=tmpdir, filename="test_doc")
            self.assertTrue(os.path.exists(path))
            self.assertTrue(path.endswith(".docx"))

    def test_docx_contains_text(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_docx(SAMPLE_TEXT, output_dir=tmpdir, filename="test_doc")
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("1890", full_text)
            self.assertIn("[illegible]", full_text)
            self.assertIn("seven children", full_text)

    def test_title_page(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_docx(
                SAMPLE_TEXT,
                output_dir=tmpdir,
                filename="titled",
                title="Family History Notebook",
            )
            doc = Document(path)
            self.assertIn("Family History Notebook", doc.paragraphs[0].text)

    def test_metadata_on_title_page(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_docx(
                SAMPLE_TEXT,
                output_dir=tmpdir,
                filename="meta",
                metadata={"pages": 8, "source_files": 2, "tokens_used": 12345},
            )
            doc = Document(path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Pages transcribed: 8", full_text)
            self.assertIn("Source files: 2", full_text)
            self.assertIn("12,345", full_text)


class TestExportRawTxt(unittest.TestCase):

    def test_creates_raw_file(self):
        """Raw export should create a .txt file with page markers."""
        raw_results = [
            {"sequence": 1, "source_file": "test.pdf", "source_page": 1,
             "text": "Page one text", "status": "ok"},
            {"sequence": 2, "source_file": "test.pdf", "source_page": 2,
             "text": "Page two text", "status": "ok"},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_raw_txt(raw_results, output_dir=tmpdir, filename="raw_out")
            self.assertTrue(os.path.exists(path))
            with open(path, encoding="utf-8") as f:
                content = f.read()
            self.assertIn("--- PAGE 1:", content)
            self.assertIn("Page one text", content)
            self.assertIn("--- PAGE 2:", content)
            self.assertIn("Page two text", content)

    def test_skips_non_ok_pages(self):
        """Raw export should skip pages with error/empty status."""
        raw_results = [
            {"sequence": 1, "source_file": "a.jpg", "source_page": 1,
             "text": "Good page", "status": "ok"},
            {"sequence": 2, "source_file": "b.jpg", "source_page": 1,
             "text": "", "status": "error"},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            path = export_raw_txt(raw_results, output_dir=tmpdir, filename="filtered")
            with open(path, encoding="utf-8") as f:
                content = f.read()
            self.assertIn("Good page", content)
            self.assertNotIn("PAGE 2", content)


class TestExportAll(unittest.TestCase):

    def test_creates_both_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            paths = export_all(SAMPLE_TEXT, output_dir=tmpdir, filename="both")
            self.assertTrue(os.path.exists(paths["txt"]))
            self.assertTrue(os.path.exists(paths["docx"]))
            self.assertNotIn("raw", paths)

    def test_creates_all_three_with_raw(self):
        """When raw_results provided, export_all should create raw .txt too."""
        raw_results = [
            {"sequence": 1, "source_file": "test.pdf", "source_page": 1,
             "text": "Raw page text", "status": "ok"},
        ]
        with tempfile.TemporaryDirectory() as tmpdir:
            paths = export_all(
                SAMPLE_TEXT, output_dir=tmpdir, filename="all3",
                raw_results=raw_results,
            )
            self.assertTrue(os.path.exists(paths["txt"]))
            self.assertTrue(os.path.exists(paths["docx"]))
            self.assertIn("raw", paths)
            self.assertTrue(os.path.exists(paths["raw"]))


if __name__ == "__main__":
    unittest.main()
